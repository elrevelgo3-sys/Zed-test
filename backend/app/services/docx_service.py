"""
DOCX Service - Generates Word documents from extracted PDF structure.

This service is responsible for:
- Creating DOCX documents that preserve the original PDF layout
- Handling text formatting (fonts, sizes, colors, bold, italic)
- Creating tables with proper structure
- Embedding images
- Managing page layout and margins
"""

import io
import os
from dataclasses import dataclass
from typing import Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


@dataclass
class DOCXSettings:
    """Settings for DOCX generation."""

    # Page settings
    page_width_inches: float = 8.5  # Letter size
    page_height_inches: float = 11.0
    margin_top_inches: float = 1.0
    margin_bottom_inches: float = 1.0
    margin_left_inches: float = 1.0
    margin_right_inches: float = 1.0

    # Default font settings
    default_font_name: str = "Arial"
    default_font_size: int = 11

    # Heading sizes
    heading_1_size: int = 18
    heading_2_size: int = 14
    heading_3_size: int = 12

    # Table settings
    table_style: str = "Table Grid"

    # Image settings
    max_image_width_inches: float = 6.5
    image_quality: int = 95


class DOCXService:
    """
    Service for generating DOCX documents from structured content.

    Converts the extracted document structure (from PDFService + OCR)
    into properly formatted Word documents.
    """

    def __init__(self, settings: Optional[DOCXSettings] = None):
        """
        Initialize the DOCX service.

        Args:
            settings: Optional custom settings for DOCX generation
        """
        self.settings = settings or DOCXSettings()

    async def generate(
        self,
        structure,  # DocumentStructure from pdf_service
        output_dir: str,
        filename: str,
        preserve_layout: bool = True,
    ) -> str:
        """
        Generate a DOCX document from the extracted structure.

        Args:
            structure: DocumentStructure with pages and elements
            output_dir: Directory to save the output file
            filename: Output filename (should end with .docx)
            preserve_layout: Whether to try preserving original layout

        Returns:
            Path to the generated DOCX file
        """
        # Create output directory if needed
        os.makedirs(output_dir, exist_ok=True)

        # Create document
        doc = Document()

        # Set up page layout
        self._setup_page_layout(doc)

        # Process each page
        for page_idx, page in enumerate(structure.pages):
            # Add page break between pages (except first)
            if page_idx > 0:
                doc.add_page_break()

            # Process elements on the page
            await self._process_page(doc, page, preserve_layout)

        # Save the document
        output_path = os.path.join(output_dir, filename)
        doc.save(output_path)

        return output_path

    def _setup_page_layout(self, doc: Document):
        """Configure page layout settings."""
        section = doc.sections[0]

        # Set page size
        section.page_width = Inches(self.settings.page_width_inches)
        section.page_height = Inches(self.settings.page_height_inches)

        # Set margins
        section.top_margin = Inches(self.settings.margin_top_inches)
        section.bottom_margin = Inches(self.settings.margin_bottom_inches)
        section.left_margin = Inches(self.settings.margin_left_inches)
        section.right_margin = Inches(self.settings.margin_right_inches)

    async def _process_page(self, doc: Document, page, preserve_layout: bool):
        """Process all elements on a page."""
        for element in page.elements:
            await self._add_element(doc, element, page.width, page.height)

    async def _add_element(
        self,
        doc: Document,
        element,  # ExtractedElement
        page_width: float,
        page_height: float,
    ):
        """Add a single element to the document."""
        element_type = element.type

        if element_type in ("paragraph", "p"):
            self._add_paragraph(doc, element)

        elif element_type in ("heading_1", "h1"):
            self._add_heading(doc, element, level=1)

        elif element_type in ("heading_2", "h2"):
            self._add_heading(doc, element, level=2)

        elif element_type in ("heading_3", "h3"):
            self._add_heading(doc, element, level=3)

        elif element_type in ("list_item", "li"):
            self._add_list_item(doc, element)

        elif element_type in ("table", "tbl"):
            self._add_table(doc, element)

        elif element_type in ("image", "img"):
            await self._add_image(doc, element)

        else:
            # Default: treat as paragraph
            self._add_paragraph(doc, element)

    def _add_paragraph(self, doc: Document, element):
        """Add a paragraph element."""
        para = doc.add_paragraph()
        run = para.add_run(element.content)

        # Apply styles
        self._apply_text_style(run, element.style)

        # Apply alignment
        alignment = element.style.get("alignment", "left") if element.style else "left"
        para.alignment = self._get_alignment(alignment)

    def _add_heading(self, doc: Document, element, level: int):
        """Add a heading element."""
        # Map level to Word heading style
        heading_style = f"Heading {level}"

        para = doc.add_paragraph(element.content, style=heading_style)

        # Apply additional formatting if needed
        if para.runs:
            run = para.runs[0]
            self._apply_text_style(run, element.style, is_heading=True, level=level)

    def _add_list_item(self, doc: Document, element):
        """Add a list item element."""
        content = element.content

        # Remove bullet/number prefix if present
        content = self._clean_list_content(content)

        # Detect if it's a numbered list or bullet list
        original = element.content.strip()
        if original and original[0].isdigit():
            # Numbered list
            para = doc.add_paragraph(content, style="List Number")
        else:
            # Bullet list
            para = doc.add_paragraph(content, style="List Bullet")

        # Apply text style
        if para.runs:
            self._apply_text_style(para.runs[0], element.style)

    def _add_table(self, doc: Document, element):
        """Add a table element."""
        # Get table data from element metadata
        table_data = element.metadata.get("rows", []) if element.metadata else []

        if not table_data:
            # If no table data, add placeholder
            doc.add_paragraph("[TABLE]")
            return

        # Determine table dimensions
        num_rows = len(table_data)
        num_cols = max(len(row) for row in table_data) if table_data else 1

        # Create table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = self.settings.table_style

        # Populate cells
        for row_idx, row_data in enumerate(table_data):
            row = table.rows[row_idx]
            for col_idx, cell_content in enumerate(row_data):
                if col_idx < len(row.cells):
                    cell = row.cells[col_idx]
                    cell.text = str(cell_content) if cell_content else ""

                    # Format cell text
                    if cell.paragraphs:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(self.settings.default_font_size)
                                run.font.name = self.settings.default_font_name

        # Add some spacing after table
        doc.add_paragraph()

    async def _add_image(self, doc: Document, element):
        """Add an image element."""
        # Check if image data is available
        image_data = element.metadata.get("image_data") if element.metadata else None
        image_base64 = (
            element.metadata.get("image_base64") if element.metadata else None
        )

        if image_data:
            # Image data as bytes
            image_stream = io.BytesIO(image_data)
            doc.add_picture(
                image_stream, width=Inches(self.settings.max_image_width_inches)
            )
        elif image_base64:
            # Image as base64 string
            import base64

            image_bytes = base64.b64decode(image_base64)
            image_stream = io.BytesIO(image_bytes)
            doc.add_picture(
                image_stream, width=Inches(self.settings.max_image_width_inches)
            )
        else:
            # Placeholder for image
            para = doc.add_paragraph("[IMAGE]")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _apply_text_style(
        self,
        run,
        style: dict,
        is_heading: bool = False,
        level: int = 0,
    ):
        """Apply text styling to a run."""
        if not style:
            style = {}

        # Font name
        font_name = style.get("font_name", self.settings.default_font_name)
        run.font.name = font_name

        # Set font for East Asian text as well
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

        # Font size
        if is_heading:
            sizes = {
                1: self.settings.heading_1_size,
                2: self.settings.heading_2_size,
                3: self.settings.heading_3_size,
            }
            font_size = sizes.get(level, self.settings.default_font_size)
        else:
            font_size = style.get("font_size", self.settings.default_font_size)

        run.font.size = Pt(font_size)

        # Bold
        if style.get("bold"):
            run.font.bold = True

        # Italic
        if style.get("italic"):
            run.font.italic = True

        # Underline
        if style.get("underline"):
            run.font.underline = True

        # Color
        color = style.get("color")
        if color and color.startswith("#"):
            try:
                # Parse hex color
                hex_color = color.lstrip("#")
                r = int(hex_color[0:2], 16)
                g = int(hex_color[2:4], 16)
                b = int(hex_color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            except (ValueError, IndexError):
                pass  # Invalid color, skip

    def _get_alignment(self, alignment: str) -> WD_ALIGN_PARAGRAPH:
        """Convert alignment string to WD_ALIGN_PARAGRAPH enum."""
        alignments = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        return alignments.get(alignment.lower(), WD_ALIGN_PARAGRAPH.LEFT)

    def _clean_list_content(self, content: str) -> str:
        """Remove bullet or number prefix from list item content."""
        content = content.strip()

        # Remove common bullet characters
        bullets = ["•", "-", "*", "–", "◦", "○", "●", "■", "□"]
        for bullet in bullets:
            if content.startswith(bullet):
                content = content[len(bullet) :].strip()
                return content

        # Remove number prefix (e.g., "1.", "1)", "a.", "a)")
        import re

        match = re.match(r"^[\d]+[.)]\s*", content)
        if match:
            content = content[match.end() :].strip()
            return content

        match = re.match(r"^[a-zA-Z][.)]\s*", content)
        if match:
            content = content[match.end() :].strip()
            return content

        return content

    def generate_from_text(
        self,
        text: str,
        output_path: str,
        title: Optional[str] = None,
    ) -> str:
        """
        Generate a simple DOCX from plain text.

        Useful for quick conversions without structure extraction.

        Args:
            text: Plain text content
            output_path: Path to save the DOCX file
            title: Optional document title

        Returns:
            Path to the generated DOCX file
        """
        doc = Document()
        self._setup_page_layout(doc)

        # Add title if provided
        if title:
            doc.add_heading(title, level=0)

        # Split text into paragraphs and add them
        paragraphs = text.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                para = doc.add_paragraph(para_text)
                for run in para.runs:
                    run.font.name = self.settings.default_font_name
                    run.font.size = Pt(self.settings.default_font_size)

        # Ensure directory exists
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        doc.save(output_path)
        return output_path

    def generate_from_html_table(
        self,
        html: str,
        output_path: str,
    ) -> str:
        """
        Generate a DOCX containing a table from HTML.

        Args:
            html: HTML string containing a table
            output_path: Path to save the DOCX file

        Returns:
            Path to the generated DOCX file
        """
        from html.parser import HTMLParser

        # Simple HTML table parser
        class TableParser(HTMLParser):
            def __init__(self):
                super().__init__()
                self.rows = []
                self.current_row = []
                self.current_cell = ""
                self.in_cell = False

            def handle_starttag(self, tag, attrs):
                if tag in ("td", "th"):
                    self.in_cell = True
                    self.current_cell = ""
                elif tag == "tr":
                    self.current_row = []

            def handle_endtag(self, tag):
                if tag in ("td", "th"):
                    self.in_cell = False
                    self.current_row.append(self.current_cell.strip())
                elif tag == "tr":
                    if self.current_row:
                        self.rows.append(self.current_row)

            def handle_data(self, data):
                if self.in_cell:
                    self.current_cell += data

        parser = TableParser()
        parser.feed(html)

        doc = Document()
        self._setup_page_layout(doc)

        if parser.rows:
            num_rows = len(parser.rows)
            num_cols = max(len(row) for row in parser.rows)

            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = self.settings.table_style

            for row_idx, row_data in enumerate(parser.rows):
                row = table.rows[row_idx]
                for col_idx, cell_content in enumerate(row_data):
                    if col_idx < len(row.cells):
                        row.cells[col_idx].text = cell_content

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        doc.save(output_path)

        return output_path
